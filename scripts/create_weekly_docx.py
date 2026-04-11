"""
Create geopolitics weekly overview Word document.
Uses the pension sample docx as template for styling.

Usage:
  python scripts/create_weekly_docx.py                          # defaults: current month, week 1
  python scripts/create_weekly_docx.py --month 2026-04 --week 2 # specific month + week
  python scripts/create_weekly_docx.py --start 2026-04-08 --end 2026-04-14  # custom date range
"""

import argparse
import json
import os
import sys
from datetime import datetime, timedelta
from pathlib import Path

from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── Project root (auto-detect from script location) ───────────────────────
SCRIPT_DIR = Path(__file__).resolve().parent
BASE_DIR = SCRIPT_DIR.parent  # track_geopolitics/

TEMPLATE = BASE_DIR / "reports" / "topics_overview" / \
    "养老金资讯-国际养老金动态（20260328-20260402）.docx"

# ── Week boundaries ───────────────────────────────────────────────────────
WEEK_RANGES = {
    1: (1, 7),
    2: (8, 14),
    3: (15, 21),
    4: (22, 31),  # month-end handled by date filtering
}


# ── Helpers ────────────────────────────────────────────────────────────────

def to_prose(text):
    """Convert bullet-format summary_zh to flowing prose paragraph."""
    if '\n' not in text:
        return text
    parts = text.split('\n')
    result = []
    for p in parts:
        p = p.strip()
        if p.startswith('- '):
            p = p[2:]
        if p:
            result.append(p)
    # Join with proper punctuation: add period between parts if missing
    joined = ""
    for i, part in enumerate(result):
        if i > 0 and joined and joined[-1] not in '。；！？…':
            joined += "。"
        joined += part
    return joined


def read_json(topic_slug, month):
    """Read a topic's monthly JSON file."""
    path = BASE_DIR / "data" / topic_slug / f"{month}.json"
    if not path.exists():
        return {"events": []}
    with open(path, 'r', encoding='utf-8') as f:
        return json.load(f)


def filter_events(data, start, end):
    """Filter events by date range (inclusive)."""
    return [e for e in data['events'] if start <= e['date'] <= end]


# ── A-tier source ranking ─────────────────────────────────────────────────

A_TIER_DOMAINS = [
    'whitehouse.gov', 'ustr.gov', 'treasury.gov', 'ofac.treasury.gov',
    'state.gov', 'commerce.gov', 'bis.gov', 'bis.doc.gov',
    'federalregister.gov', 'eia.gov', 'energy.gov',
    'fmprc.gov.cn', 'mofcom.gov.cn', 'english.news.cn', 'news.cn',
    'ndrc.gov.cn', 'gov.cn',
    'ec.europa.eu', 'eca.europa.eu', 'consilium.europa.eu',
    'europarl.europa.eu',
    'opec.org', 'wto.org', 'imf.org', 'worldbank.org', 'iea.org',
    'un.org',
]

A_TIER_NAMES = [
    'OFAC', 'USTR', 'BIS', '白宫', 'EIA', '美国EIA', '美国商务部',
    '外交部', '商务部', '新华社', '发改委',
    'European Commission', 'European Court', 'Auditors',
    'OPEC', 'WTO', 'IMF', 'IEA', 'Federal Register',
]


def _source_rank(src):
    """Lower = more authoritative (A-tier preferred)."""
    url = src.get('url', '')
    name = src.get('name', '')
    for domain in A_TIER_DOMAINS:
        if domain in url:
            return 0
    for pat in A_TIER_NAMES:
        if pat in name:
            return 1
    return 2


def build_source_line(event):
    """Pick the most authoritative source with URL."""
    sources = event['sources']
    ranked = sorted(sources, key=_source_rank)
    src = ranked[0]
    name = src['name']
    url = src.get('url', '')
    if url:
        return f"资料来源：{name}，{url}"
    else:
        return f"资料来源：{name}，{event['date']}"


# ── Country mapping ───────────────────────────────────────────────────────

COUNTRY_MAP = [
    ("OFAC", "美国"),
    ("特朗普", "美国"),
    ("商务部修订", "美国"),
    ("商务部调整", "美国"),
    ("EIA", "美国"),
    ("安森美", "美国"),
    ("Section 232", "美国"),
    ("联邦", "美国"),
    ("欧盟", "欧盟"),
    ("欧洲审计院", "欧盟"),
    ("CBAM", "欧盟"),
    ("OPEC", "多国"),
    ("CMI", "国际"),
    ("中国外交部", "中国"),
    ("DeepSeek", "中国"),
]


def get_country(title):
    """Determine the country/entity prefix for event title."""
    for key, val in COUNTRY_MAP:
        if key in title:
            return val
    return ""


# ── Document building ─────────────────────────────────────────────────────

def get_body_style(doc):
    """Find the 'Paragraph' custom style; fall back to 'Normal'."""
    for s in doc.styles:
        if s.name == 'Paragraph':
            return 'Paragraph'
    return 'Normal'


def add_toc_field(doc):
    """Add a Word TOC field that can be updated in Word (Ctrl+A → F9)."""
    p_heading = doc.add_paragraph()
    p_heading.alignment = 1  # center
    run = p_heading.add_run("目  录")
    run.bold = True
    run.font.size = Pt(10)

    paragraph = doc.add_paragraph()

    # Begin field
    run1 = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    run1._r.append(fldChar1)

    # Field instruction
    run2 = paragraph.add_run()
    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    instrText.text = ' TOC \\o "2-3" \\h \\z \\u '
    run2._r.append(instrText)

    # Separate
    run3 = paragraph.add_run()
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'separate')
    run3._r.append(fldChar2)

    # Placeholder text
    run4 = paragraph.add_run("（请右键此处 → 更新域 → 更新整个目录）")
    run4.font.size = Pt(10)

    # End field
    run5 = paragraph.add_run()
    fldChar3 = OxmlElement('w:fldChar')
    fldChar3.set(qn('w:fldCharType'), 'end')
    run5._r.append(fldChar3)

    doc.add_paragraph("")


# ── Main ───────────────────────────────────────────────────────────────────

def main():
    parser = argparse.ArgumentParser(
        description="Generate geopolitics weekly overview Word document"
    )
    parser.add_argument('--month', default=None,
                        help='Target month (YYYY-MM). Default: current month')
    parser.add_argument('--week', type=int, default=1, choices=[1, 2, 3, 4],
                        help='Week number (1-4). Default: 1')
    parser.add_argument('--start', default=None,
                        help='Custom start date (YYYY-MM-DD), overrides --week')
    parser.add_argument('--end', default=None,
                        help='Custom end date (YYYY-MM-DD), overrides --week')
    args = parser.parse_args()

    # Determine month
    if args.month:
        month = args.month
    else:
        month = datetime.now().strftime('%Y-%m')

    year, mon = month.split('-')

    # Determine date range
    if args.start and args.end:
        start_date = args.start
        end_date = args.end
    else:
        day_start, day_end = WEEK_RANGES[args.week]
        start_date = f"{month}-{day_start:02d}"
        # Clamp end to actual month end
        if args.week == 4:
            # Find last day of month
            if int(mon) == 12:
                last_day = 31
            else:
                from calendar import monthrange
                _, last_day = monthrange(int(year), int(mon))
            end_date = f"{month}-{last_day:02d}"
        else:
            end_date = f"{month}-{day_end:02d}"

    # Derive output filename
    s = start_date.replace('-', '')
    e = end_date.replace('-', '')
    out_name = f"地缘政治追踪-国际地缘政治动态（{s}-{e}）.docx"
    out_dir = BASE_DIR / "reports" / "topics_overview" / month
    out_path = out_dir / out_name

    print("=" * 60)
    print("Creating geopolitics weekly overview docx")
    print(f"  Month: {month}  |  Range: {start_date} ~ {end_date}")
    print(f"  Output: {out_path.name}")
    print("=" * 60)

    # 1. Open template
    if not TEMPLATE.exists():
        print(f"ERROR: Template not found: {TEMPLATE}")
        sys.exit(1)
    doc = Document(str(TEMPLATE))

    # 2. Clear body content (keep sectPr)
    body = doc.element.body
    children_to_remove = []
    for child in body:
        tag = child.tag.split('}')[-1] if '}' in child.tag else child.tag
        if tag == 'sectPr':
            continue
        children_to_remove.append(child)
    for child in children_to_remove:
        body.remove(child)
    print(f"  Cleared {len(children_to_remove)} template elements")

    body_style = get_body_style(doc)

    # 3. Load event data
    topics = [
        ("topic1-sanctions", "制裁与经济管制", "一"),
        ("topic2-trade-industrial", "贸易与产业政策", "二"),
        ("topic3-supply-resources", "能源安全与资源角力", "三"),
        ("topic4-tech-digital", "技术竞争与规则制定", "四"),
    ]

    all_events = {}
    for slug, name, num in topics:
        data = read_json(slug, month)
        events = filter_events(data, start_date, end_date)
        all_events[slug] = events
        print(f"  {name}: {len(events)} events")

    total = sum(len(v) for v in all_events.values())
    print(f"  Total: {total} events")

    if total == 0:
        print("\nNo events found in date range. Exiting.")
        sys.exit(0)

    # 4. Build document
    # Title
    title_text = f"地缘政治追踪-国际地缘政治动态（{s}-{e}）"
    doc.add_paragraph(title_text, style='Heading 1')

    # TOC
    add_toc_field(doc)

    # Sections
    n = 0
    for slug, section_name, num_zh in topics:
        events = all_events[slug]
        doc.add_paragraph(f"{num_zh}、{section_name}", style='Heading 2')

        if not events:
            doc.add_paragraph("（本周暂无事件）", style=body_style)
            continue

        for event in events:
            n += 1
            country = get_country(event['title_zh'])
            date_parts = event['date'].split('-')
            date_str = f"{date_parts[1]}-{date_parts[2]}"
            title = f"{date_str}|{country}|{event['title_zh']}"

            doc.add_paragraph(title, style='Heading 3')
            doc.add_paragraph(to_prose(event['summary_zh']), style=body_style)
            doc.add_paragraph(build_source_line(event), style=body_style)

    # 5. Save
    out_dir.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))
    print(f"\n  ✓ Saved: {out_path}")
    print(f"  Events: {n}")
    print(f"  提示：Word中右键目录 → 更新域 → 更新整个目录")
    print("=" * 60)


if __name__ == "__main__":
    main()
